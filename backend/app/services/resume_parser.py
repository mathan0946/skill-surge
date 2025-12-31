"""
Robust Resume Parser Service for SkillSurge
Supports PDF, DOCX, and TXT files with intelligent section detection
"""
import re
import io
from typing import Optional, Dict, List, Tuple
from dataclasses import dataclass


@dataclass
class ParsedResume:
    """Structured resume data after parsing"""
    raw_text: str
    sections: Dict[str, str]
    contact_info: Dict[str, str]
    detected_skills: List[str]
    detected_experience: List[Dict]
    detected_education: List[Dict]
    detected_projects: List[Dict]
    detected_certifications: List[str]
    parsing_confidence: float  # 0-1 score of parsing quality


class ResumeParser:
    """
    Multi-format resume parser with intelligent section detection
    """
    
    # Common section headers in resumes
    SECTION_PATTERNS = {
        'experience': r'(?i)^(?:work\s*)?experience|employment\s*history|professional\s*experience|work\s*history',
        'education': r'(?i)^education|academic\s*background|qualifications|academic\s*credentials',
        'skills': r'(?i)^(?:technical\s*)?skills|competencies|technologies|tech\s*stack|expertise|proficiencies',
        'projects': r'(?i)^projects|personal\s*projects|portfolio|notable\s*projects|key\s*projects',
        'certifications': r'(?i)^certifications?|licenses?|credentials|professional\s*certifications?',
        'achievements': r'(?i)^achievements?|accomplishments?|awards?|honors?|recognition',
        'summary': r'(?i)^(?:professional\s*)?summary|objective|profile|about\s*me|career\s*objective',
        'publications': r'(?i)^publications?|papers?|research',
        'languages': r'(?i)^languages?|spoken\s*languages',
        'interests': r'(?i)^interests?|hobbies|activities',
        'references': r'(?i)^references?',
    }
    
    # Common technical skills to detect
    TECH_SKILLS = [
        # Programming Languages
        'Python', 'JavaScript', 'TypeScript', 'Java', 'C++', 'C#', 'Go', 'Rust', 'Ruby', 'PHP', 
        'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl', 'Objective-C', 'Dart', 'Lua',
        # Frontend
        'React', 'Vue', 'Angular', 'Next.js', 'Nuxt', 'Svelte', 'HTML', 'CSS', 'SASS', 'SCSS',
        'Tailwind', 'Bootstrap', 'jQuery', 'Redux', 'MobX', 'Webpack', 'Vite', 'Babel',
        # Backend
        'Node.js', 'Express', 'FastAPI', 'Django', 'Flask', 'Spring Boot', 'Rails', 'Laravel',
        'ASP.NET', 'GraphQL', 'REST', 'gRPC', 'Microservices',
        # Databases
        'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Elasticsearch', 'SQLite', 'Oracle', 
        'SQL Server', 'DynamoDB', 'Cassandra', 'Neo4j', 'Supabase', 'Firebase',
        # Cloud & DevOps
        'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI', 'GitHub Actions',
        'Terraform', 'Ansible', 'Linux', 'Nginx', 'Apache',
        # AI/ML
        'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'Pandas', 'NumPy', 'OpenAI', 'LangChain',
        'Machine Learning', 'Deep Learning', 'NLP', 'Computer Vision', 'LLM',
        # Mobile
        'React Native', 'Flutter', 'iOS', 'Android', 'SwiftUI', 'Jetpack Compose',
        # Tools
        'Git', 'Jira', 'Confluence', 'Figma', 'Postman', 'VS Code', 'IntelliJ',
        # Methodologies
        'Agile', 'Scrum', 'CI/CD', 'TDD', 'BDD', 'DevOps', 'Microservices Architecture',
    ]
    
    # Email pattern
    EMAIL_PATTERN = r'[\w\.-]+@[\w\.-]+\.\w+'
    # Phone pattern
    PHONE_PATTERN = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    # LinkedIn pattern
    LINKEDIN_PATTERN = r'linkedin\.com/in/[\w-]+'
    # GitHub pattern
    GITHUB_PATTERN = r'github\.com/[\w-]+'
    # URL pattern
    URL_PATTERN = r'https?://[^\s<>"{}|\\^`\[\]]+'
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'doc', 'txt', 'rtf']
    
    async def parse(self, file_content: bytes, filename: str) -> ParsedResume:
        """
        Main entry point for parsing a resume file
        """
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if ext == 'pdf':
            raw_text = await self._extract_pdf_text(file_content)
        elif ext in ['docx', 'doc']:
            raw_text = await self._extract_docx_text(file_content)
        elif ext == 'txt':
            raw_text = self._extract_txt_text(file_content)
        else:
            # Try PDF first, then plain text
            raw_text = await self._extract_pdf_text(file_content)
            if not raw_text.strip():
                raw_text = self._extract_txt_text(file_content)
        
        # Clean and normalize the text
        cleaned_text = self._clean_text(raw_text)
        
        # Parse structured data
        sections = self._detect_sections(cleaned_text)
        contact_info = self._extract_contact_info(cleaned_text)
        skills = self._detect_skills(cleaned_text)
        experience = self._parse_experience(sections.get('experience', ''))
        education = self._parse_education(sections.get('education', ''))
        projects = self._parse_projects(sections.get('projects', ''))
        certifications = self._parse_certifications(sections.get('certifications', ''))
        
        # Calculate parsing confidence
        confidence = self._calculate_confidence(sections, skills, experience)
        
        return ParsedResume(
            raw_text=cleaned_text,
            sections=sections,
            contact_info=contact_info,
            detected_skills=skills,
            detected_experience=experience,
            detected_education=education,
            detected_projects=projects,
            detected_certifications=certifications,
            parsing_confidence=confidence
        )
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """
        Extract text from PDF using multiple methods for robustness
        """
        text = ""
        
        # Method 1: PyPDF2 (primary)
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
        
        # Method 2: pdfplumber (better for complex layouts)
        if not text.strip():
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                pass  # pdfplumber not installed
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
        
        # Method 3: pdfminer (most robust for text extraction)
        if not text.strip():
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                text = pdfminer_extract(io.BytesIO(content))
            except ImportError:
                pass  # pdfminer not installed
            except Exception as e:
                print(f"pdfminer extraction failed: {e}")
        
        # Method 4: PyMuPDF/fitz (fastest and handles many edge cases)
        if not text.strip():
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except ImportError:
                pass  # PyMuPDF not installed
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")
        
        return text
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """
        Extract text from DOCX files
        """
        text = ""
        
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
                    
        except ImportError:
            print("python-docx not installed, trying alternative method")
            # Fallback: try to extract as XML
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    doc_xml = zf.read('word/document.xml')
                    tree = ET.fromstring(doc_xml)
                    # Extract all text nodes
                    for elem in tree.iter():
                        if elem.text:
                            text += elem.text + " "
            except Exception as e:
                print(f"DOCX XML extraction failed: {e}")
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
        
        return text
    
    def _extract_txt_text(self, content: bytes) -> str:
        """
        Extract text from plain text files
        """
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, AttributeError):
                continue
        
        # Last resort: decode with errors ignored
        return content.decode('utf-8', errors='ignore')
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize resume text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace but preserve structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive spaces within lines
            line = re.sub(r' {3,}', '  ', line)
            # Remove control characters except newlines
            line = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', line)
            # Strip trailing/leading whitespace
            line = line.strip()
            
            if line:
                cleaned_lines.append(line)
        
        # Join lines, collapsing multiple blank lines
        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Detect and extract resume sections using pattern matching
        """
        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if this line is a section header
            section_found = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped):
                    section_found = section_name
                    break
            
            # Also check for common header formats: ALL CAPS, ending with colon
            if not section_found and line_stripped:
                # ALL CAPS line with 2+ words is likely a header
                if line_stripped.isupper() and len(line_stripped.split()) <= 4:
                    for section_name, pattern in self.SECTION_PATTERNS.items():
                        if re.search(pattern, line_stripped):
                            section_found = section_name
                            break
                # Line ending with colon
                elif line_stripped.endswith(':'):
                    header_text = line_stripped[:-1]
                    for section_name, pattern in self.SECTION_PATTERNS.items():
                        if re.search(pattern, header_text):
                            section_found = section_name
                            break
            
            if section_found:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract contact information from resume
        """
        contact = {}
        
        # Email
        email_match = re.search(self.EMAIL_PATTERN, text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_match = re.search(self.PHONE_PATTERN, text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group()
        
        # GitHub
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        if github_match:
            contact['github'] = github_match.group()
        
        # Name (usually first non-empty line)
        lines = text.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Skip if it looks like contact info
            if line and not re.search(self.EMAIL_PATTERN, line) and \
               not re.search(self.PHONE_PATTERN, line) and \
               not re.search(self.URL_PATTERN, line):
                # Check if it looks like a name (2-4 words, capitalized)
                words = line.split()
                if 1 <= len(words) <= 4:
                    if all(w[0].isupper() for w in words if w):
                        contact['name'] = line
                        break
        
        return contact
    
    def _detect_skills(self, text: str) -> List[str]:
        """
        Detect technical skills mentioned in the resume
        """
        found_skills = []
        text_lower = text.lower()
        
        for skill in self.TECH_SKILLS:
            # Use word boundary matching for accuracy
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text, re.IGNORECASE):
                found_skills.append(skill)
        
        # Also look for skills in parentheses or after colons (common skill list format)
        skill_list_patterns = [
            r'(?:skills?|technologies?|tech stack|stack)[\s:]+([^\n]+)',
            r'\(([^)]+)\)',  # Parenthetical mentions
        ]
        
        for pattern in skill_list_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Split by common delimiters
                potential_skills = re.split(r'[,;|/•·]', match)
                for ps in potential_skills:
                    ps = ps.strip()
                    if ps and len(ps) < 30:  # Reasonable skill name length
                        # Check if it matches any known skill
                        for known in self.TECH_SKILLS:
                            if known.lower() in ps.lower():
                                if known not in found_skills:
                                    found_skills.append(known)
        
        return found_skills
    
    def _parse_experience(self, experience_text: str) -> List[Dict]:
        """
        Parse work experience section into structured data
        """
        experiences = []
        if not experience_text:
            return experiences
        
        # Pattern for date ranges
        date_pattern = r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*)?(?:\d{4}|\d{1,2}/\d{2,4})'
        date_range_pattern = rf'({date_pattern})\s*[-–—to]+\s*({date_pattern}|Present|Current|Now)'
        
        # Split by likely job entries (look for patterns like company names, dates)
        lines = experience_text.split('\n')
        current_job = {}
        current_bullets = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for date range (indicates new job)
            date_match = re.search(date_range_pattern, line, re.IGNORECASE)
            
            # Check if line looks like a job title or company
            looks_like_header = (
                line.isupper() or 
                '|' in line or 
                date_match or
                re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*[-–—|,]', line)
            )
            
            if looks_like_header and (current_job or date_match):
                # Save previous job
                if current_job:
                    current_job['highlights'] = current_bullets
                    experiences.append(current_job)
                
                current_job = {
                    'company': '',
                    'role': '',
                    'duration': '',
                    'highlights': []
                }
                current_bullets = []
                
                # Extract dates
                if date_match:
                    current_job['startDate'] = date_match.group(1)
                    current_job['endDate'] = date_match.group(2)
                    current_job['duration'] = f"{date_match.group(1)} - {date_match.group(2)}"
                    # Remove date from line for further parsing
                    line = re.sub(date_range_pattern, '', line).strip()
                
                # Try to split company and role
                if '|' in line:
                    parts = line.split('|')
                    current_job['company'] = parts[0].strip()
                    if len(parts) > 1:
                        current_job['role'] = parts[1].strip()
                elif ',' in line:
                    parts = line.split(',')
                    current_job['role'] = parts[0].strip()
                    if len(parts) > 1:
                        current_job['company'] = parts[1].strip()
                else:
                    current_job['role'] = line
                    
            elif line.startswith(('•', '-', '●', '○', '*', '▪', '►')):
                # Bullet point
                bullet = line.lstrip('•-●○*▪► ').strip()
                if bullet:
                    current_bullets.append(bullet)
            elif current_job:
                # Add as continuation of company/role info
                if not current_job.get('company'):
                    current_job['company'] = line
                elif not current_job.get('role'):
                    current_job['role'] = line
        
        # Save last job
        if current_job:
            current_job['highlights'] = current_bullets
            experiences.append(current_job)
        
        return experiences
    
    def _parse_education(self, education_text: str) -> List[Dict]:
        """
        Parse education section into structured data
        """
        education = []
        if not education_text:
            return education
        
        # Common degree patterns
        degree_pattern = r'(?:Bachelor|Master|Ph\.?D\.?|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|M\.?B\.?A\.?|B\.?E\.?|B\.?Tech|M\.?Tech|Associate)[\'s]*\.?\s*(?:of|in)?\s*(?:Science|Arts|Engineering|Business|Computer|Technology)?[^,\n]*'
        
        lines = education_text.split('\n')
        current_edu = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for degree
            degree_match = re.search(degree_pattern, line, re.IGNORECASE)
            
            # Check for year
            year_match = re.search(r'(19|20)\d{2}', line)
            
            # Check for GPA
            gpa_match = re.search(r'(?:GPA|CGPA)[\s:]*(\d+\.?\d*)', line, re.IGNORECASE)
            
            if degree_match or (line.isupper() and len(line.split()) <= 6):
                # Save previous education
                if current_edu:
                    education.append(current_edu)
                
                current_edu = {
                    'degree': degree_match.group() if degree_match else line,
                    'institution': '',
                    'year': year_match.group() if year_match else '',
                    'gpa': gpa_match.group(1) if gpa_match else ''
                }
            elif current_edu and not current_edu.get('institution'):
                # This line is likely the institution
                current_edu['institution'] = line
        
        # Save last education
        if current_edu:
            education.append(current_edu)
        
        return education
    
    def _parse_projects(self, projects_text: str) -> List[Dict]:
        """
        Parse projects section into structured data
        """
        projects = []
        if not projects_text:
            return projects
        
        lines = projects_text.split('\n')
        current_project = {}
        current_description = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Project title usually starts with capital, may have link or date
            looks_like_title = (
                (line[0].isupper() if line else False) and 
                not line.startswith(('•', '-', '●', '○', '*')) and
                len(line.split()) <= 10
            )
            
            if looks_like_title:
                # Save previous project
                if current_project:
                    current_project['description'] = ' '.join(current_description)
                    current_project['technologies'] = self._detect_skills(' '.join(current_description))
                    projects.append(current_project)
                
                current_project = {
                    'name': line,
                    'description': '',
                    'technologies': [],
                    'link': ''
                }
                current_description = []
                
                # Check for GitHub link
                github_match = re.search(self.URL_PATTERN, line)
                if github_match:
                    current_project['link'] = github_match.group()
                    current_project['name'] = re.sub(self.URL_PATTERN, '', line).strip()
            else:
                # Description line
                bullet = line.lstrip('•-●○*▪► ').strip()
                if bullet:
                    current_description.append(bullet)
        
        # Save last project
        if current_project:
            current_project['description'] = ' '.join(current_description)
            current_project['technologies'] = self._detect_skills(' '.join(current_description))
            projects.append(current_project)
        
        return projects
    
    def _parse_certifications(self, cert_text: str) -> List[str]:
        """
        Parse certifications into a list
        """
        certifications = []
        if not cert_text:
            return certifications
        
        lines = cert_text.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                # Remove bullet points
                cert = line.lstrip('•-●○*▪► ').strip()
                if cert and len(cert) > 3:
                    certifications.append(cert)
        
        return certifications
    
    def _calculate_confidence(self, sections: Dict, skills: List, experience: List) -> float:
        """
        Calculate a confidence score for the parsing quality
        """
        score = 0.0
        
        # Has structured sections
        if len(sections) > 2:
            score += 0.3
        elif len(sections) > 0:
            score += 0.15
        
        # Has skills detected
        if len(skills) > 10:
            score += 0.25
        elif len(skills) > 5:
            score += 0.15
        elif len(skills) > 0:
            score += 0.1
        
        # Has experience parsed
        if len(experience) > 2:
            score += 0.25
        elif len(experience) > 0:
            score += 0.15
        
        # Has specific sections
        if 'experience' in sections:
            score += 0.1
        if 'education' in sections:
            score += 0.05
        if 'projects' in sections:
            score += 0.05
        
        return min(score, 1.0)


# Singleton instance
resume_parser = ResumeParser()


async def parse_resume_file(content: bytes, filename: str) -> ParsedResume:
    """
    Parse a resume file and return structured data
    """
    return await resume_parser.parse(content, filename)


def get_enhanced_prompt_context(parsed: ParsedResume) -> str:
    """
    Generate an enhanced context string for the OpenAI prompt
    """
    context_parts = []
    
    # Contact info
    if parsed.contact_info:
        context_parts.append(f"CANDIDATE: {parsed.contact_info.get('name', 'Unknown')}")
    
    # Pre-detected skills
    if parsed.detected_skills:
        context_parts.append(f"\nPRE-DETECTED SKILLS: {', '.join(parsed.detected_skills)}")
    
    # Sections
    for section_name, content in parsed.sections.items():
        if content.strip():
            context_parts.append(f"\n=== {section_name.upper()} ===\n{content[:3000]}")
    
    # If we have parsed experience, include it
    if parsed.detected_experience:
        context_parts.append("\n=== PARSED EXPERIENCE ===")
        for exp in parsed.detected_experience[:5]:
            context_parts.append(f"- {exp.get('role', 'Role')} at {exp.get('company', 'Company')} ({exp.get('duration', '')})")
    
    # Parsing confidence
    context_parts.append(f"\n[Parsing Confidence: {parsed.parsing_confidence:.0%}]")
    
    return '\n'.join(context_parts)
